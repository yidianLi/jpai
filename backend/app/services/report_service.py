"""报告生成服务：Word报告 + 数据图表"""
import json
from datetime import datetime, date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from ..database import AiSessionLocal
from .analysis_service import AnalysisService
from .warning_service import WarningService
from .idle_service import IdleService
from ..models.report import AiReport

# 设置中文字体
plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False


class ReportService:
    def __init__(self):
        self.db = AiSessionLocal()
        self.analysis = AnalysisService()
        self.warning = WarningService()
        self.idle = IdleService()

    def close(self):
        self.db.close()
        self.analysis.close()
        self.warning.close()
        self.idle.close()

    def generate_monthly_report(self, year, month, user="system"):
        """生成月度资产管理报告"""
        period = f"{year}-{month:02d}"
        overview = self.analysis.get_overview()
        class_dist = self.analysis.get_class_distribution()
        dept_rank = self.analysis.get_dept_ranking()
        warnings = self.warning.get_warning_list(status=0, size=100)
        idle_stats = self.idle.get_idle_stats()

        snapshot_time = datetime.now().isoformat()
        content = {
            "period": period, "overview": overview, "class_distribution": class_dist,
            "department_ranking": dept_rank, "warnings": warnings, "idle_stats": idle_stats,
            "generate_time": snapshot_time,
            "snapshot": {
                "captured_at": snapshot_time,
                "data_cutoff": overview.get("data_cutoff"),
                "rules_version": "report-v1",
                "ai_used": False,
                "metric_definitions": overview.get("metric_definitions", {}),
                "source_tables": ["ai_asset", "ai_asset_transfer", "ai_warning", "ai_idle_pool"]
            }
        }

        # 生成Word
        doc = Document()
        # 标题
        title = doc.add_heading(f"{period} 固定资产管理分析报告", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()

        # 一、资产概况
        doc.add_heading("一、资产概况", level=1)
        p = doc.add_paragraph()
        p.add_run(f"截至报告期末，本单位固定资产总数 {overview['total_count']} 台/件，")
        p.add_run(f"资产原值 {overview['total_value']:,.2f} 元，")
        p.add_run(f"账面净值 {overview['current_value']:,.2f} 元。")
        p.add_run(f"其中在用 {overview['in_use_count']} 台，闲置 {overview['idle_count']} 台，")
        p.add_run(f"已报废 {overview['scrap_count']} 台。")
        p.add_run(f"闲置率 {overview['idle_rate']}%，账实相符率 {overview['match_rate']}%。")

        # 二、分类分布
        doc.add_heading("二、资产分类分布", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = "资产类别"
        hdr[1].text = "数量(台/件)"
        hdr[2].text = "原值(元)"
        for item in class_dist[:10]:
            row = table.add_row().cells
            row[0].text = item["name"]
            row[1].text = str(item["count"])
            row[2].text = f"{item['value']:,.2f}"

        # 三、部门排名
        doc.add_heading("三、部门资产排名", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = "部门"
        hdr[1].text = "资产数量"
        hdr[2].text = "资产原值(元)"
        hdr[3].text = "闲置率"
        for item in dept_rank[:10]:
            row = table.add_row().cells
            row[0].text = item["dept"]
            row[1].text = str(item["count"])
            row[2].text = f"{item['value']:,.2f}"
            row[3].text = f"{item['idle_rate']}%"

        # 四、预警情况
        doc.add_heading("四、预警情况", level=1)
        doc.add_paragraph(f"当前未处理预警共 {warnings['total']} 条，其中：")
        red = sum(1 for w in warnings["list"] if w["level"] == 1)
        yellow = sum(1 for w in warnings["list"] if w["level"] == 2)
        doc.add_paragraph(f"红色预警（紧急）{red} 条，黄色预警（提醒）{yellow} 条。", style='List Bullet')
        if warnings["list"]:
            doc.add_paragraph("重点预警事项：")
            for w in warnings["list"][:5]:
                doc.add_paragraph(w["content"], style='List Number')

        # 五、闲置资产
        doc.add_heading("五、闲置资产与盘活建议", level=1)
        doc.add_paragraph(
            f"当前闲置资产 {idle_stats['idle_count']} 台，估算价值 {idle_stats['idle_value']:,.2f} 元。"
            f"建议对闲置超过180天的资产进行内部调拨或公开处置，提高资产使用效率。"
        )

        # 六、下月工作建议
        doc.add_heading("六、下月工作建议", level=1)
        suggestions = [
            "对到期预警资产进行逐一核查，及时办理报废或续用手续",
            "组织闲置资产内部调剂会，优先在单位内部调配使用",
            "开展账实不符资产专项清理，确保账实相符率持续提升",
            "完善资产入库验收流程，从源头保证数据质量",
        ]
        for s in suggestions:
            doc.add_paragraph(s, style='List Number')

        file_path = f"reports/monthly_report_{period}.docx"
        import os
        os.makedirs("reports", exist_ok=True)
        doc.save(file_path)

        # 保存记录
        record = AiReport(
            report_type="monthly", title=f"{period}固定资产管理分析报告",
            # SQLAlchemy 聚合金额可能是 Decimal，标准 JSON 编码器无法直接保存。
            # 以字符串保留精度，避免报告生成成功后在落库阶段失败。
            period=period, content=json.dumps(content, ensure_ascii=False, default=str),
            file_path=file_path, create_user=user, create_time=datetime.now()
        )
        self.db.add(record)
        self.db.commit()
        return {"id": record.id, "title": record.title, "file_path": file_path, "period": period, "ai_used": False,
                "message": "报告已按规则统计生成；当前数据未调用大模型润色。"}

    def get_report_list(self, report_type=None, page=1, size=20):
        q = self.db.query(AiReport)
        if report_type:
            q = q.filter(AiReport.report_type == report_type)
        total = q.count()
        rows = q.order_by(AiReport.create_time.desc()).offset((page-1)*size).limit(size).all()
        return {
            "total": total, "page": page, "size": size,
            "list": [{"id": r.id, "type": r.report_type, "title": r.title,
                      "period": r.period, "file_path": r.file_path,
                      "create_time": str(r.create_time),
                      "snapshot": self._snapshot_metadata(r.content)} for r in rows]
        }

    @staticmethod
    def _snapshot_metadata(content):
        try:
            data = json.loads(content or "{}")
            return data.get("snapshot", {"ai_used": False, "rules_version": "legacy"})
        except (TypeError, ValueError):
            return {"ai_used": False, "rules_version": "legacy"}
